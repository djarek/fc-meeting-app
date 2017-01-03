# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

import apps.council.models as council_models


def generate_attendance_list(meeting):
    document = Document()

    # Setting margins
    sections = document.sections
    for section in sections:
        section.left_margin = Cm(5)
        section.top_margin = Cm(5)

    # Date of creation document
    paragraph = document.add_paragraph(unicode(generate_place_date(), 'utf-8'))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # break
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    run.add_break()

    # title
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'{} - lista obecności'.format(meeting))
    run.bold = True

    # table
    table = document.add_table(rows=1, cols=4)
    table.style = 'TableGrid'

    # table header
    table_header_cols = (u'Stopień/tytuł naukowy', u'Imię', u'Nazwisko',
                         u'Podpis')
    row = table.rows[0]
    for num, text in enumerate(table_header_cols):
        paragraph = row.cells[num].paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True

    # table content
    invited_list = council_models.Invited.objects.filter(meeting=meeting)
    for invited in invited_list:
        row = table.add_row()
        row.cells[0].text = invited.person.get_title_verbose()
        row.cells[1].text = invited.person.first_name
        row.cells[2].text = invited.person.last_name
        row.cells[3].text = u''

    return document


def generate_invitation_letter(meeting):
    document = Document()

    # Setting margins
    sections = document.sections
    for section in sections:
        section.left_margin = Cm(5)
        section.top_margin = Cm(5)

    # Date of creation document
    paragraph = document.add_paragraph(unicode(generate_place_date(), 'utf-8'))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # break
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    run.add_break()

    # introduction
    paragraph = document.add_paragraph(
        u'    Uprzejmie zapraszam na posiedzenie Rady Wydziału Elektroniki'
        u' Politechniki Wrocławskiej, które odbędzie się w dniu {} o godzinie'
        u' {}, {}'.format(
            meeting.date.strftime('%d.%m.%Y r.'),
            meeting.date.strftime('%H:%M'),
            meeting.place
        )
    )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'Porządek obrad:')
    run.bold = True

    # Points list
    points = council_models.Point.objects.filter(meeting=meeting).\
        order_by('number')
    for point in points:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(u'{}.   {}'.format(point.number, point.title))
        run.bold = True
        run = paragraph.add_run(u'\nRodzaj sprawy: {}'.format(point.category))
        run = paragraph.add_run(u'\nOpis: {}'.format(point.description))

    return document


def generate_place_date():
    now = datetime.now()
    sDateTime = "Wrocław, dnia " + now.strftime("%d.%m.%Y") + " r."
    return sDateTime
